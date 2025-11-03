"""
文件上传和预览API视图
支持 PDF, Word, Excel, TXT 文件的上传和预览
"""
import os
import json
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

from django.conf import settings
from django.http import JsonResponse, FileResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from loguru import logger


# 支持的文件类型和扩展名
SUPPORTED_EXTENSIONS = {
    'pdf': ['.pdf'],
    'word': ['.doc', '.docx'],
    'excel': ['.xls', '.xlsx', '.csv'],
    'text': ['.txt']
}

# 所有支持的扩展名列表
ALL_EXTENSIONS = []
for exts in SUPPORTED_EXTENSIONS.values():
    ALL_EXTENSIONS.extend(exts)


def get_file_type(filename: str) -> str:
    """根据文件名获取文件类型"""
    ext = Path(filename).suffix.lower()
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if ext in extensions:
            return file_type
    return 'unknown'


@csrf_exempt
@api_view(['POST'])
def upload_file(request):
    """
    上传文件接口
    支持 PDF, Word, Excel, TXT 文件
    
    POST /api/files/upload/
    """
    try:
        # 检查是否有文件上传
        if 'file' not in request.FILES:
            return Response({
                'success': False,
                'error': '没有上传文件'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        filename = uploaded_file.name
        file_ext = Path(filename).suffix.lower()
        
        # 验证文件类型
        if file_ext not in ALL_EXTENSIONS:
            return Response({
                'success': False,
                'error': f'不支持的文件类型。支持的格式：{", ".join(ALL_EXTENSIONS)}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        logger.info(f"收到文件上传请求: {filename}")
        
        # 创建上传目录
        upload_dir = Path(settings.BASE_DIR) / 'media' / 'uploads'
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        # 生成唯一的文件名（使用时间戳）
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_stem = Path(filename).stem
        unique_filename = f"{file_stem}_{timestamp}{file_ext}"
        file_path = upload_dir / unique_filename
        
        # 保存文件
        with open(file_path, 'wb+') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        logger.info(f"文件保存成功: {file_path}")
        
        # 获取文件信息
        file_size = os.path.getsize(file_path)
        file_type = get_file_type(filename)
        
        # 生成文件ID（相对于 media/uploads 的路径）
        file_id = unique_filename
        
        return Response({
            'success': True,
            'file_id': file_id,
            'filename': filename,
            'original_filename': filename,
            'file_type': file_type,
            'file_size': file_size,
            'upload_time': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"上传文件处理失败: {traceback.format_exc()}")
        return Response({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def preview_file(request, file_id: str):
    """
    文件预览接口
    根据文件类型返回不同的预览内容
    
    GET /api/files/preview/<file_id>/
    """
    try:
        upload_dir = Path(settings.BASE_DIR) / 'media' / 'uploads'
        file_path = upload_dir / file_id
        
        if not file_path.exists():
            return Response({
                'success': False,
                'error': '文件不存在'
            }, status=status.HTTP_404_NOT_FOUND)
        
        file_type = get_file_type(file_id)
        
        # PDF 文件：返回文件供前端 PDF.js 渲染
        if file_type == 'pdf':
            return FileResponse(
                open(file_path, 'rb'),
                content_type='application/pdf',
                as_attachment=False
            )
        
        # TXT 文件：返回文本内容
        elif file_type == 'text':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return Response({
                'success': True,
                'file_type': 'text',
                'content': content
            })
        
        # Word 文件：转换为 HTML
        elif file_type == 'word':
            html_content = convert_word_to_html(file_path)
            return Response({
                'success': True,
                'file_type': 'word',
                'content': html_content
            })
        
        # Excel 文件：转换为 JSON 数据
        elif file_type == 'excel':
            excel_data = convert_excel_to_json(file_path)
            return Response({
                'success': True,
                'file_type': 'excel',
                'data': excel_data
            })
        
        else:
            return Response({
                'success': False,
                'error': '不支持的文件类型'
            }, status=status.HTTP_400_BAD_REQUEST)
        
    except Exception as e:
        logger.error(f"预览文件失败: {traceback.format_exc()}")
        return Response({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def download_file(request, file_id: str):
    """
    文件下载接口
    
    GET /api/files/download/<file_id>/
    """
    try:
        upload_dir = Path(settings.BASE_DIR) / 'media' / 'uploads'
        file_path = upload_dir / file_id
        
        if not file_path.exists():
            return Response({
                'success': False,
                'error': '文件不存在'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # 返回文件
        response = FileResponse(
            open(file_path, 'rb'),
            as_attachment=True,
            filename=file_path.name
        )
        return response
        
    except Exception as e:
        logger.error(f"下载文件失败: {traceback.format_exc()}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def convert_word_to_html(file_path: Path) -> str:
    """
    将 Word 文档转换为 HTML
    """
    try:
        # 尝试使用 python-docx 处理 .docx 文件
        if file_path.suffix.lower() == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                
                html_parts = ['<div class="word-content">']
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        html_parts.append(f'<p>{para.text}</p>')
                
                # 处理表格
                for table in doc.tables:
                    html_parts.append('<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">')
                    for row in table.rows:
                        html_parts.append('<tr>')
                        for cell in row.cells:
                            html_parts.append(f'<td style="padding: 8px; border: 1px solid #ddd;">{cell.text}</td>')
                        html_parts.append('</tr>')
                    html_parts.append('</table>')
                
                html_parts.append('</div>')
                return ''.join(html_parts)
                
            except ImportError:
                logger.warning("python-docx 未安装，无法解析 .docx 文件")
                return '<p>无法预览 Word 文档。请安装 python-docx 库。</p>'
        
        # .doc 文件需要其他库（如 antiword 或 LibreOffice）
        else:
            return '<p>暂不支持 .doc 格式预览，请使用 .docx 格式。</p>'
            
    except Exception as e:
        logger.error(f"转换 Word 文档失败: {e}")
        return f'<p>预览失败: {str(e)}</p>'


def convert_excel_to_json(file_path: Path) -> Dict[str, Any]:
    """
    将 Excel/CSV 文件转换为 JSON 数据
    """
    try:
        import pandas as pd
        import numpy as np
        
        def clean_value(val):
            """将值转换为JSON可序列化的Python原生类型"""
            if pd.isna(val):
                return None
            elif isinstance(val, (np.integer, np.int64, np.int32)):
                return int(val)
            elif isinstance(val, (np.floating, np.float64, np.float32)):
                if np.isinf(val) or np.isnan(val):
                    return None
                return float(val)
            elif isinstance(val, np.bool_):
                return bool(val)
            elif isinstance(val, (np.ndarray, list)):
                return [clean_value(v) for v in val]
            else:
                return val
        
        def clean_data_for_json(df):
            """清理数据使其可以JSON序列化"""
            # 替换 inf 和 -inf 为 None
            df = df.replace([np.inf, -np.inf], None)
            # 将 NaN 替换为 None
            df = df.where(pd.notnull(df), None)
            
            # 转换所有值为Python原生类型
            data = []
            for row in df.values:
                cleaned_row = [clean_value(val) for val in row]
                data.append(cleaned_row)
            
            return df.columns.tolist(), data
        
        # 判断文件类型
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.csv':
            # CSV 文件只有一个 sheet
            df = pd.read_csv(file_path)
            columns, data = clean_data_for_json(df)
            
            sheets_data = {
                'Sheet1': {
                    'columns': columns,
                    'data': data,
                    'rows': len(data),
                    'cols': len(columns)
                }
            }
            return {
                'sheets': ['Sheet1'],
                'data': sheets_data
            }
        else:
            # Excel 文件可能有多个 sheet
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                columns, data = clean_data_for_json(df)
                
                # 转换为字典列表
                sheets_data[sheet_name] = {
                    'columns': columns,
                    'data': data,
                    'rows': len(data),
                    'cols': len(columns)
                }
            
            return {
                'sheets': list(sheets_data.keys()),
                'data': sheets_data
            }
        
    except ImportError:
        logger.error("pandas 或 openpyxl 未安装，无法解析 Excel 文件")
        return {
            'error': '无法预览 Excel 文档。请安装 pandas 和 openpyxl 库。'
        }
    except Exception as e:
        logger.error(f"转换 Excel 文档失败: {e}")
        return {
            'error': f'预览失败: {str(e)}'
        }


@api_view(['DELETE'])
def delete_file(request, file_id: str):
    """
    删除文件接口
    
    DELETE /api/files/<file_id>/
    """
    try:
        upload_dir = Path(settings.BASE_DIR) / 'media' / 'uploads'
        file_path = upload_dir / file_id
        
        if not file_path.exists():
            return Response({
                'success': False,
                'error': '文件不存在'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # 删除文件
        os.remove(file_path)
        logger.info(f"文件已删除: {file_path}")
        
        return Response({
            'success': True,
            'message': '文件删除成功'
        })
        
    except Exception as e:
        logger.error(f"删除文件失败: {traceback.format_exc()}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
